import io
import os
import json
import time
from typing import List, Dict, Any
from pathlib import Path

import pandas as pd
import streamlit as st
from dotenv import load_dotenv
from docx import Document  # NEW

from page_extractor import extract_pagewise, extract_image, build_excel_sheets, DEFAULT_MODEL

# ----------------- env -----------------
load_dotenv()
load_dotenv(dotenv_path=Path(__file__).parent.parent / ".env")

# ----------------- UI helpers -----------------
def _progress_renderer():
    stage_line = st.empty()
    bar = st.progress(0.0)
    page_times: List[float] = []
    state = {"done_pages": 0, "total_pages_est": 0}

    def cb(ev: Dict[str, Any]):
        stage = ev.get("stage")
        i, n = int(ev.get("i", 0)), int(ev.get("n", 0))
        if n:
            state["total_pages_est"] = n
        labels = {
            "page_start": "Starting page",
            "ocrmypdf": "OCRmyPDF (whole file prepass)",
            "pdftotext": "Extracting text layer",
            "preprocess": "Preprocessing (denoise/deskew/threshold)",
            "ocr": "OCR (best-of passes)",
            "ocr_try": f"OCR try (langs={ev.get('langs')}, psm={ev.get('psm')})",
            "tables": "Extracting tables",
            "ai": "Extracting AI fields",
            "page_done": "Finishing page",
        }
        if stage in labels:
            info = labels[stage]
            if i and n:
                info = f"{info} — page {i}/{n}"
            stage_line.write(f"**Stage:** {info}")

        if stage == "page_done":
            sec = float(ev.get("sec", 0.0))
            if sec > 0:
                page_times.append(sec)
            state["done_pages"] += 1

        # progress bar
        if state["total_pages_est"]:
            frac = state["done_pages"] / max(1, state["total_pages_est"])
            bar.progress(min(1.0, frac))

    return cb

def _download_buttons(doc_json: dict, excel_sheets: Dict[str, pd.DataFrame], docx_bytes: bytes | None):
    # JSON
    json_bytes = json.dumps(doc_json, ensure_ascii=False, indent=2).encode("utf-8")
    st.download_button("⬇️ Download JSON", data=json_bytes, file_name="doc.json", mime="application/json")

    # Excel
    out = io.BytesIO()
    with pd.ExcelWriter(out, engine="xlsxwriter") as w:
        for name, df in excel_sheets.items():
            # Excel sheet name limit 31
            df.to_excel(w, index=False, sheet_name=name[:31])
    st.download_button("⬇️ Download Excel", data=out.getvalue(),
                       file_name="doc.xlsx",
                       mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    # Word
    if docx_bytes:
        st.download_button("⬇️ Download Word (.docx)", data=docx_bytes,
                           file_name="doc.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

def _build_word(doc) -> bytes:
    """Compose a simple, clean .docx from extracted content."""
    d = Document()
    d.add_heading("AI-Readable Extraction", level=0)
    meta_p = d.add_paragraph()
    meta_p.add_run(f"doc_id: {doc.doc_id}\n")
    if getattr(doc, "doc_type", None):
        meta_p.add_run(f"doc_type: {doc.doc_type}\n")

    # Pages
    for p in doc.pages:
        d.add_heading(f"Page {p.page_index}", level=1)
        # Meta
        if p.lang or p.ocr:
            d.add_paragraph(f"lang: {p.lang or '-'} | ocr: {p.ocr}")
        # Text
        if p.text and p.text.strip():
            d.add_paragraph(p.text.strip())

        # Key/Value fields
        if p.kv_fields:
            d.add_paragraph("AI Fields:")
            tbl = d.add_table(rows=1, cols=5)
            hdr = tbl.rows[0].cells
            hdr[0].text = "Key"
            hdr[1].text = "Value"
            hdr[2].text = "Type"
            hdr[3].text = "Confidence"
            hdr[4].text = "Page"
            for k, fv in p.kv_fields.items():
                row = tbl.add_row().cells
                row[0].text = str(k)
                row[1].text = "" if fv.value is None else str(fv.value)
                row[2].text = fv.type
                row[3].text = "" if fv.confidence is None else str(fv.confidence)
                row[4].text = "" if fv.page is None else str(fv.page)

        # Tables detected from page
        if p.tables:
            for t in p.tables:
                d.add_paragraph(f"Table: {t.name}")
                cols = [c.name for c in t.columns]
                tbl = d.add_table(rows=1, cols=len(cols))
                hdr = tbl.rows[0].cells
                for j, name in enumerate(cols):
                    hdr[j].text = str(name)
                for r in t.rows:
                    row = tbl.add_row().cells
                    for j in range(len(cols)):
                        row[j].text = "" if j >= len(r) or r[j] is None else str(r[j])

    # Normalized sections (if any)
    if getattr(doc, "normalized", None):
        if "first_schedule_items" in doc.normalized:
            d.add_heading("Normalized: First Schedule Items", level=1)
            items = doc.normalized["first_schedule_items"]
            if items:
                keys = ["item_no", "hs_code", "description", "rate", "appendix", "page"]
                tbl = d.add_table(rows=1, cols=len(keys))
                for j, k in enumerate(keys):
                    tbl.rows[0].cells[j].text = k
                for it in items:
                    row = tbl.add_row().cells
                    for j, k in enumerate(keys):
                        row[j].text = "" if it.get(k) is None else str(it.get(k))
        if "second_schedule_items" in doc.normalized:
            d.add_heading("Normalized: Second Schedule Items", level=1)
            items = doc.normalized["second_schedule_items"]
            if items:
                tbl = d.add_table(rows=1, cols=3)
                hdr = ["page", "description", "doc_id"]
                for j, k in enumerate(hdr):
                    tbl.rows[0].cells[j].text = k
                for it in items:
                    row = tbl.add_row().cells
                    row[0].text = "" if it.get("page") is None else str(it.get("page"))
                    row[1].text = "" if it.get("description") is None else str(it.get("description"))
                    row[2].text = "" if it.get("doc_id") is None else str(it.get("doc_id"))

    bio = io.BytesIO()
    d.save(bio)
    return bio.getvalue()

# ----------------- App -----------------
st.set_page_config(page_title="PDF → AI-Readable (PP-OCR-like)", layout="wide")
st.title("📄 PDF → CSV/JSON (with optional PP-OCR-like det+rec)")

with st.sidebar:
    st.header("Settings")

    aggressive_ocr = st.checkbox(
        "Aggressive preprocess (denoise/deskew)",
        value=True,
        help="OpenCV denoise, adaptive threshold, deskew."
    )
    try_ocrmypdf = st.checkbox(
        "Try OCRmyPDF pre-pass",
        value=True,
        help="If the text layer is weak, run OCRmyPDF once for the whole file."
    )
    speed_mode = st.checkbox(
        "Speed mode (fast)",
        value=True,
        help="Use fewer Tesseract passes."
    )
    # PP-OCR-like toggle
    use_ppocr_like = st.checkbox(
        "Experimental: PP-OCR-like (det+rec lines)",
        value=False,
        help="Uses easyocr if installed; otherwise falls back to Tesseract line boxes."
    )
    # NEW: GPU toggle for EasyOCR
    use_gpu = st.checkbox(
        "Use CUDA for EasyOCR (if available)",
        value=False,
        help="Sets EASYOCR_GPU=1. Requires NVIDIA driver + CUDA-enabled PyTorch."
    )
    os.environ["EASYOCR_GPU"] = "1" if use_gpu else "0"

    # Show PyTorch CUDA status if installed
    try:
        import torch  # noqa
        st.caption(f"PyTorch CUDA: {'available' if torch.cuda.is_available() else 'not available'}")
        if getattr(torch, 'cuda', None) and torch.cuda.is_available():
            st.caption(f"GPU: {torch.cuda.get_device_name(0)} | CUDA {torch.version.cuda}")
    except Exception:
        st.caption("PyTorch not installed or import failed (only needed for EasyOCR).")

    st.divider()
    call_ai = st.checkbox("Call OpenAI for fields/summary", value=False)
    api_key = st.text_input("OPENAI_API_KEY", type="password", value=os.getenv("OPENAI_API_KEY") or "")
    model = st.text_input("Model", value=DEFAULT_MODEL)

    fields_raw = st.text_area(
        "Fields to extract (comma-separated)",
        value="HS Code, Description, Rate",
        help="Used only if Call OpenAI is enabled."
    )
    fields = [f.strip() for f in fields_raw.split(",") if f.strip()] if call_ai else None

st.write("Upload a **PDF** or an **image** (PNG/JPG).")

up = st.file_uploader("Choose file", type=["pdf", "png", "jpg", "jpeg"], accept_multiple_files=False)

if up:
    fb = up.read()
    ext = (up.name.split(".")[-1] or "").lower()
    cb = _progress_renderer()

    try:
        if ext == "pdf":
            doc = extract_pagewise(
                fb, up.name,
                api_key=(api_key if call_ai else ""),
                model=model,
                fields=fields,
                aggressive_ocr=aggressive_ocr,
                try_ocrmypdf=try_ocrmypdf,
                speed_mode=speed_mode,
                progress_cb=cb,
                use_ppocr_like=use_ppocr_like,   # PP-OCR-like det+rec
            )
        else:
            doc = extract_image(
                fb, up.name,
                api_key=(api_key if call_ai else ""),
                model=model,
                fields=fields,
                aggressive_ocr=aggressive_ocr,
                speed_mode=speed_mode,
                progress_cb=cb,
                use_ppocr_like=use_ppocr_like,   # PP-OCR-like det+rec
            )
    except Exception as e:
        st.error(f"Extraction error: {e}")
        st.stop()

    # ---- Results ----
    st.success("Done!")

    # per-page preview
    for p in doc.pages:
        with st.expander(f"Page {p.page_index}"):
            col1, col2 = st.columns([2, 1])
            with col1:
                st.text_area("Text", p.text, height=220, key=f"text_{doc.doc_id}_{p.page_index}")
            with col2:
                st.write("Meta")
                st.json({"lang": p.lang, "ocr_used": p.ocr})

            # lines table
            if getattr(p, "lines", None):
                line_rows = []
                for L in p.lines:
                    poly = L.poly if isinstance(L.poly, list) else []
                    poly = (poly + [[None, None]] * 4)[:4]
                    line_rows.append({
                        "x1": poly[0][0], "y1": poly[0][1],
                        "x2": poly[1][0], "y2": poly[1][1],
                        "x3": poly[2][0], "y3": poly[2][1],
                        "x4": poly[3][0], "y4": poly[3][1],
                        "text": L.text, "conf": L.conf,
                    })
                st.caption("Detected lines")
                st.dataframe(pd.DataFrame(line_rows), use_container_width=True)

            # kv fields (AI)
            if p.kv_fields:
                rows = [{"key": k, **v.model_dump()} for k, v in p.kv_fields.items()]
                st.caption("AI Key/Value Fields")
                st.dataframe(pd.DataFrame(rows), use_container_width=True)

            # tables
            if p.tables:
                st.caption("Tables")
                for t in p.tables:
                    df = pd.DataFrame(t.rows, columns=[c.name for c in t.columns])
                    st.write(f"**{t.name}**")
                    st.dataframe(df, use_container_width=True)

    # doc meta/normalized
    with st.expander("Document Meta & Normalized Outputs", expanded=False):
        st.json({
            "doc_id": doc.doc_id,
            "doc_type": doc.doc_type,
            "rules": doc.rules,
            "meta": doc.meta,
            "normalized_keys": list((doc.normalized or {}).keys())
        })

    # downloads (JSON, Excel, Word)
    sheets = build_excel_sheets(doc)
    docx_bytes = _build_word(doc)
    _download_buttons(json.loads(doc.model_dump_json()), sheets, docx_bytes)

# friendly footer
st.caption("Tip: Set env `OCR_LANGS=chi_sim+chi_tra+eng+msa`. Toggle 'Use CUDA' if you installed a CUDA-enabled PyTorch. EasyOCR uses EASYOCR_GPU=1.")
